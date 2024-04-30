import json
import sys
# Markdown
import markdown
import os
# pyperclip
import pyperclip
import subprocess

from datetime import datetime, date

# odfpy
from odf.opendocument import OpenDocumentText
from odf.style import Style, TextProperties
from odf.text import H, P

# python-docx 
from docx import Document

from InstallPyQt6 import installPyQt6
installPyQt6()

# PyQt6
from PyQt6 import uic
from PyQt6.QtWidgets import QApplication, QMainWindow, QFileDialog
from PyQt6.QtGui import QIcon

from dialogs.about import About
from dialogs.messagebox import Messagebox

class MyWindow(QMainWindow):
    version = '1.9'
    # Global configuration
    config = {}

    # Dictionary with all notes
    notes = {}

    # Current note
    currentNote = {}

    currentNoteNumber = 0

    # Content of current MD file
    content =''

    def __init__(self):
        """Initializes the application.

        The GUI is built and signals and slots are connected.
        """
        m = '__init__'
        self.__log__(m)

        super(MyWindow, self).__init__()
        self.ui = uic.loadUi("./PersonalNotes.ui", self)
        self.readConfig()
        self.openNotes()
        self.refreshList()
        self.setCurrentNote(0)

        # Add icons to buttons
        self.ui.pushButton_Add.setIcon(QIcon('./icons/add.svg'))
        self.ui.pushButton_Del.setIcon(QIcon('./icons/del.svg'))
        self.ui.pushButton_Copy.setIcon(QIcon('./icons/copy.svg'))
        self.ui.pushButton_Mail.setIcon(QIcon('./icons/mail.svg'))

        # Slots for menu items
        self.ui.action_Open_Folder.triggered.connect(self.onOpenFolder)
        self.ui.action_save_as_MD.triggered.connect(self.onSaveAsMD)
        self.ui.action_save_as_HTML.triggered.connect(self.onSaveAsHTML)
        self.ui.action_save_as_ODT.triggered.connect(self.onSaveAsODT)
        self.ui.action_save_as_DOCX.triggered.connect(self.onSaveAsDOCX)
        self.ui.action_Exit.triggered.connect(self.onExit)
        self.ui.action_About.triggered.connect(self.onAbout)

        #Slots for buttons
        self.ui.listWidget.itemClicked.connect(self.onItemClicked)
        self.ui.pushButton_Add.clicked.connect(self.onButtonAdd)
        self.ui.pushButton_Del.clicked.connect(self.onButtonDel)
        self.ui.pushButton_Copy.clicked.connect(self.onButtonCopy)
        self.ui.pushButton_Mail.clicked.connect(self.onButtonMail)

        QApplication.instance().focusChanged.connect(self.onFocusChanged)
    

    def __log__(self, message):
        """Saves the call of the methods in a log file.

        The time and method name are recorded in the log file.
        Logging is only active if the value "Logging": "True" is set in config.json.
        """
        timestamp = datetime.now()
        s = str(timestamp) + ': ' + message
        
        try:
            if self.config['Logging'] == 'True':
                with open('./logfile.log', 'a') as f:
                    f.write(s + '\n')
        except:
            pass


    def readConfig(self):
        """Reads the configuration.

        Reads the config.json and saves the content in the sef.config dictionary.
        """
        m = 'readConfig'
        self.__log__(m)

        try:
            with open('./config.json', 'r') as f:
                self.config = json.load(f)
        except:
            self.config = {
                                "NotesFolder": "./Notes/",
                                "NotesFilename": "notes.json",
                                "CommentFolder": "./Comments/",
                                "ChecklistFolder": "./Checklists/",
                                "PictureFolder": "./Pictures/",
                                "Logging": "False"
            }
             
            self.saveConfig()


    def saveConfig(self):
        """Saves the configuration

        Saves the dictionary self.config in the config.json file.
        """
        m = 'saveConfig'
        self.__log__(m)

        with open('./config.json', 'w') as f:
            json.dump(self.config, f)


    def openNotes(self):
        """Opens the Notes.

        Reads the notes.json file into the dictionary self.notes.
        """
        m = 'openNotes'
        self.__log__(m)

        file_name = self.config['NotesFolder'] + self.config['NotesFilename'] 
        try:
            with open(file_name, 'r') as f:
                self.notes = json.load(f)

                # No creation date in file
                for i in range(len(self.notes['item'])):
                    item = self.notes['item'][i]
               
                    if not 'date_creation' in item:
                        # Add key 'date_creation' to the dictionary
                        self.notes['item'][i]['date_creation'] = self.notes['item'][i]['date']

                    if not 'tags' in item:
                        # Add empty key 'tags' to the dictionary
                        self.notes['item'][i]['tags'] = ''
                
        except:
            Messagebox.info('Notes file not found. A new file is created.')
            self.notes = {
                            "LastNoteID": 1,
                            "LastNote": 1,
                            "LastComment": 1,
                            "LastChecklist": 1,
                            "item": []
            }

            self.saveNotes()


    def saveNotes(self):
        """Saves the notes.

        Saves the dictionary self.notes in the notes.json file.
        """
        m = 'saveNotes'
        self.__log__(m)
        
        file_name = self.config['NotesFolder'] + self.config['NotesFilename']
        with open(file_name, 'w') as f:
            json.dump(self.notes, f)


    def refreshList(self):
        """Renews the list of notes.
        """
        m = 'refreshList'
        self.__log__(m)

        itemList = []
        for i in range(len(self.notes['item'])):
            s = self.notes['item'][i]['name']
            if s[0] == '#':
                s = s.lstrip('#')
                s = s.lstrip()

            if len(s) > 40:
                s = s[:40] + ' ...\n'

            s = s + self.notes['item'][i]['date'] + '\n'
            itemList.append(s)

        self.ui.listWidget.clear()
        self.ui.listWidget.addItems(itemList)
        self.ui.listWidget.setCurrentRow(0)


    def setCurrentNote(self, itemNumber):
        """Activates the selected note.

        Reads the item specified with itemNumber from the note list and saves it in self.currentNote.

        Keyword arguments:
        itemNumber -- Number of the item that is to be set active.
        """
        m = 'setCurrentNote'
        self.__log__(m)

        self.currentNote = self.notes['item'][itemNumber]
        self.displayContent(self.currentNote['note'])


    def displayContent(self, file):
        """Displays the content of the Markdown file.

        Reads the Markdown file into self.content and displays it in textEdit.

        Keyword arguments:
        file -- Name of the Markdown file.
        """
        m = 'displayContent'
        self.__log__(m)

        fullFileName = self.config['NotesFolder'] + file
        # TODO #65 Error handling if the Markdown file not found
        with open(fullFileName, 'r') as f:
            self.content = f.read()

        # Convert Markdown into HTML
        contentHTML = markdown.markdown(self.content)

        self.ui.textEdit.setHtml(contentHTML)
        self.ui.textEdit.setReadOnly(True)


    def onOpenFolder(self):
        """Opens a folder.

        Handles the "Open Folder" menu item.
        Opens a folder with notes. The notes in this folder are read in.
        """
        m = 'onOpenFolder'
        self.__log__(m)

        directory = QFileDialog.getExistingDirectory(self, "Open Folder")
        self.config['NotesFolder'] = directory + '/'
        
        self.saveConfig()
        self.openNotes()
        self.refreshList()
        self.setCurrentNote(0)


    def onSaveAsMD(self):
        """Saves the Markdown file.

        Handles the "Save as Markdown" menu item.
        Saves the Markdown file in a folder.
        """
        m = 'onSaveAsMD'
        self.__log__(m)

        directory = QFileDialog.getExistingDirectory(self, "Open Folder")
        file_name = directory + '/' + self.currentNote['note']
        with open(file_name, 'w') as f:
            f.write(self.content)


    def onSaveAsHTML(self):
        """Saves the content as HTML

        Handles the "Save as HTML" menu item.
        Saves the content as an HTML file in a folder.
        """
        m = 'onSaveAsHTML'
        self.__log__(m)

        directory = QFileDialog.getExistingDirectory(self, "Open Folder")
        file_name = directory + '/' + self.currentNote['note']
        # Change the file extension
        file_name = file_name.rstrip('.md') + '.html'
        # convert Markdown into HTML
        html = '''
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>''' + self.currentNote['name'].rstrip('\n') +  '''</title>
</head>
<body>
''' + markdown.markdown(self.content) + '''
</body>
</html>
        '''
        with open(file_name, 'w') as f:
            f.write(html)


    def onSaveAsODT(self):
        """Saves the content as ODT

        Handles the "Save as ODT" menu item.
        Saves the content as an OpenDocument file in a folder.
        """
        m = 'onSaveAsODT'
        self.__log__(m)

        directory = QFileDialog.getExistingDirectory(self, "Open Folder")
        file_name = directory + '/' + self.currentNote['note']

        # Change the file extension
        file_name = file_name.rstrip('.md') + '.odt'

        # Create an instance of OpenDocumentText
        textDoc = OpenDocumentText()

        # Styles
        myStyle = textDoc.styles

        h1Style = Style(name="Heading 1", family="paragraph")
        h1Style.addElement(TextProperties(attributes={'fontsize':"24pt",'fontweight':"bold" }))
        myStyle.addElement(h1Style)

        h2Style = Style(name="Heading 2", family="paragraph")
        h2Style.addElement(TextProperties(attributes={'fontsize':"20pt",'fontweight':"bold" }))
        myStyle.addElement(h2Style)

        h3Style = Style(name="Heading 3", family="paragraph")
        h3Style.addElement(TextProperties(attributes={'fontsize':"16pt",'fontweight':"bold" }))
        myStyle.addElement(h3Style)

        # TODO #89 Implement a style for lists in ODT

        # Convert text into a list
        textList = self.content.split('\n')
        
        for line in textList:
            if line == '':
                # Don't write blank lines into ODT
                pass
            elif line[:3] == '###':
                # Heading 3
                line = line.lstrip('#')
                line = line.lstrip()
                h = H(outlinelevel=3, stylename=h3Style, text=str(line))
                textDoc.text.addElement(h)
            elif line[:2] == '##':
                # Heading 2
                line = line.lstrip('#')
                line = line.lstrip()
                h = H(outlinelevel=2, stylename=h2Style, text=str(line))
                textDoc.text.addElement(h)
            elif line[:1] =='#':
                # Heading 1
                line = line.lstrip('#')
                line = line.lstrip()
                h = H(outlinelevel=1, stylename=h1Style, text=str(line))
                textDoc.text.addElement(h)
            else:
                # Paragraf
                p = P(text=str(line))
                textDoc.text.addElement(p)
            
        textDoc.save(file_name)


    def onSaveAsDOCX(self):
        """Saves the content as DOCX

        Handles the "Save as DOCX" menu item.
        Saves the content as an Microsoft Word file in a folder.
        """
        m = 'onSaveAsDOCX'
        self.__log__(m)

        directory = QFileDialog.getExistingDirectory(self, "Open Folder")
        file_name = directory + '/' + self.currentNote['note']

        # Change the file extension
        file_name = file_name.rstrip('.md') + '.docx'

        # Create an instance of Document
        document = Document()

        # Convert text into a list
        textList = self.content.split('\n')

        for line in textList:
            if line == '':
                # Don't write blank lines into DOCX
                pass
            elif line[:3] == '###':
                # Heading 3
                line = line.lstrip('#')
                line = line.lstrip()
                document.add_heading(line, level=3)
            elif line[:2] == '##':
                # Heading 2
                line = line.lstrip('#')
                line = line.lstrip()
                document.add_heading(line, level=2)
            elif line[:1] =='#':
                # Heading 1
                line = line.lstrip('#')
                line = line.lstrip()
                document.add_heading(line, level=1)
            else:
                # Paragraf
                document.add_paragraph(line)

        document.save(file_name)


    def onExit(self):
        """Close the Application

        Handles the "Exit" menu item.
        """
        m = 'onExit'
        self.__log__(m)
        
        sys.exit(0)


    def onAbout(self):
        About(self.version)


    def onItemClicked(self, item):
        """Handles the click on an entry in the list.
        """
        m = 'onItemClicked'
        self.__log__(m)

        # Number of selected line
        self.currentNoteNumber = self.ui.listWidget.currentRow()
        self.setCurrentNote(self.currentNoteNumber)


    def onButtonAdd(self):
        """Adds a new entry.

        Handles the click on the Add button.
        Creates a new entry in the list and creates a new Markdown file.
        """
        m = 'onButtonAdd'
        self.__log__(m)

        self.ui.listWidget.setFocus()
        # Write the MD file
        self.notes['LastNote'] = self.notes['LastNote'] +1
        file = self.notes['LastNote']
        fullFileName = self.config['NotesFolder'] + str(file) + '.md'
        with open(fullFileName, 'w') as f:
            f.write('')
        # Add a new item into dictionary
        today = date.today()
        newItem ={
            "name": "New Note\n",
            "date": str(today),
            "note": str(file) + '.md',
            "date_creation": str(today)
        }
        self.notes['item'].insert(0, newItem)
        self.saveNotes()
        self.refreshList()


    def onButtonDel(self):
        """Deletes the current entry

        Handles the click on the Del button.
        Deletes the current entry from the list and removes the Markdown file.
        """
        m = 'onButtonDel'
        self.__log__(m)

        self.ui.listWidget.setFocus()
        # Which item is in focus?
        row = self.listWidget.currentRow()
        
        # Delete Markdown file
        fullFileName = self.config['NotesFolder'] + self.currentNote['note']

        if os.path.exists(fullFileName):
            os.remove(fullFileName)
        else:
            # TODO: #41 Replace it with a message box
            print("The file does not exist.")
        
        # Delete the item in list
        self.notes['item'].remove(self.notes['item'][row])

        self.saveNotes()
        self.refreshList()
        self.setCurrentNote(0)
        self.ui.textEdit.setReadOnly(True)


    def onButtonCopy(self):
        """Copies the content to the clipboard.

        Handles the click on the copy button.
        """
        m = 'onButtonCopy'
        self.__log__(m)

        pyperclip.copy(self.content)


    def onButtonMail(self):
        """Sends the content as an e-mail.

        Handles the click on the mail button.
        The system's e-mail program is opened and the content of the note is 
        copied to a new e-mail.
        """
        m = 'onButtonMail'
        self.__log__(m)

        address = 'email@example.com'
        l = self.content.splitlines(True)
        if len(l) > 0:
            subject = l[0]
        else:
            subject = 'New Item\n'

        if subject[0] == '#':
            subject = subject.lstrip('#')
            subject = subject.lstrip()

        body = self.content

        if sys.platform == 'darwin':
            # macOS
            subprocess.call(["open", "mailto:" + address + "?subject=" + subject + "&body=" + body])
        else:
            # Windows, Linux
            os.system("start mailto:" + address + "?subject=" + subject + "&body=" + body)


    def onFocusChanged(self, oldWidget, nowWidget):
        """Changing the focus.

        This method is executed when the focus is changed in the app.

        Keyword arguments:
        oldWidget -- Widget that was focused before the change
        nowWidget -- Widget which is now focused
        """
        m = 'onFocusChanged'
        self.__log__(m)

        if self.ui.textEdit is nowWidget:
            # QTextEdit is in focus
            # TODO: #32 Change the background color

            # Make QTextEdit editable
            self.ui.textEdit.setReadOnly(False)
            self.ui.textEdit.setPlainText(self.content)

        if self.ui.textEdit is oldWidget:
            # QTextEdit is no longer in focus
            # TODO: #32 Change the background color
            
            self.content = self.ui.textEdit.toPlainText()
            
            # Save the file
            file_name = self.config['NotesFolder'] + self.currentNote['note']
            with open(file_name, 'w') as f:
                f.write(self.content)

            # Display the content as HTML
            self.displayContent(self.currentNote['note'])
            
            # Change the date and move it on the 1st place
            today = date.today()
            newItem = {}
            l = self.content.splitlines(True)
            if len(l) > 0:
                name = l[0]
            else:
                name = 'New Item\n'
            
            newItem['name'] = name
            newItem['date'] = str(today)
            newItem['note'] = self.currentNote['note']
            newItem['date_creation'] = self.currentNote['date_creation']
            
            del(self.notes['item'][self.currentNoteNumber])
            self.notes['item'].insert(0, newItem)
            
            self.saveNotes()
            self.refreshList()
                    

def window():
    app = QApplication(sys.argv)
    with open('./styles.qss', 'r') as f:
        style = f.read()
        app.setStyleSheet(style)
    win = MyWindow()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
	window()
